import os
import shutil
import openpyxl
import xml.etree.ElementTree as ET
import docx
import string
from docx.enum.section import WD_ORIENT

startingCol = 142
startingRow = 7412


def main():
    try:
        if os.path.isfile('Result.docx'):
            os.remove('Result.docx')
        shutil.copy('src1','Result.docx')
    except OSError:
        print('Result.docx error')
        exit(1)
    try:
        wb = openpyxl.load_workbook('target.xlsx')
        ws = wb['VPHiddenSheet']
    except OSError:
        print("target.xlsx not found")
        exit(1)

    count = 0
    while ws.cell(row=startingRow + count, column=startingCol).value is not None:
        count += 1

    # Check for errors before continuing

    # Get raw xml string from LEME configured ELN form
    rawXML = ''
    for x in range(0, count):
        rawXML += ws.cell(row=startingRow + x, column=startingCol).value

    # Add cell settings from XML
    tree = ET.fromstring(rawXML)
    cell_settings = tree.find('cellproperties')

    # Sort cell names alphabetically
    data = []
    for cell_property in cell_settings:
        key = cell_property.findtext('.//name')
        data.append((key, cell_property))

    data.sort()

    new_data = []
    current_item = ''
    store_item = ()
    len_data = len(data)
    for item in data:
        if current_item == '' or current_item == -1:
            new_data.append(item)
            if item[0][item[0].rfind('_')+1:].isnumeric():
                current_item = item[0][:item[0].rfind('_')]
        else:
            if item[0][:item[0].rfind('_')] == current_item:
                store_item = item
            else:
                last_new_data = new_data.pop()
                last_new_data[1].find('.//name').text = last_new_data[1].findtext('.//name') + " - " + store_item[1].findtext('.//name')
                last_new_data[1].find('.//address').text = last_new_data[1].findtext('.//address') + ":" + store_item[1].findtext('.//address')
                new_data.append(last_new_data)
                new_data.append(item)
                if item[0][item[0].rfind('_') + 1:].isnumeric():
                    current_item = item[0][:item[0].rfind('_')]
                else:
                    current_item = -1

    cell_settings[:] = [item[-1] for item in new_data]

    # Write to Word 2010 Document
    # document = docx.Document()
    document = docx.Document('Result.docx')
    # document.sections[-1].orientation = WD_ORIENT.LANDSCAPE

    # Create Cell Settings table with header
    table = document.add_table(1, 4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sheet'
    hdr_cells[1].text = 'Address'
    hdr_cells[2].text = 'Name'
    hdr_cells[3].text = 'Cell Settings'

    for cellProperty in cell_settings:
        cell = cellProperty.find('cell')
        row_cells = table.add_row().cells
        row_cells[0].text = cell.find('sheet').text
        row_cells[1].text = cell.find('address').text
        row_cells[2].text = cell.find('name').text
        for prop in cellProperty.findall('property'):
            for x in prop.iter():
                if x.tag == 'guid' or x is prop:
                    continue
                elif x.tag == 'sds':
                    for y in list(x):
                        if y.tag == 'sdsParameters':
                            tempTable = row_cells[3].add_table(1, 3)
                            tempHdr_cells = tempTable.rows[0].cells
                            tempHdr_cells[0].text = 'Name'
                            tempHdr_cells[1].text = 'Output'
                            tempHdr_cells[2].text = 'Unit'
                            for z in list(y):
                                tempRow_cells = tempTable.add_row().cells
                                tempRow_cells[0].text = z.find('name').text
                                try:
                                    tempRow_cells[1].text = z.find('output').find('cell').find('name').text
                                except:
                                    pass
                                try:
                                    tempRow_cells[2].text = z.find('unitoutput').find('cell').find('name').text
                                except:
                                    pass
                                print(row_cells[3].text)
                        else:
                            for z in y.iter():
                                if z.text is not None and z.tag != 'guid' and z is not y:
                                    row_cells[3].text += string.capwords(z.tag) + ': ' + str(z.text) + '\n'
                    continue
                elif x.text is not None:
                    row_cells[3].text += string.capwords(x.tag) + ': ' + str(x.text) + '\n'
                else:
                    row_cells[3].text += string.capwords(x.tag) + '\n'
            row_cells[3].text += '\n'

    document.save('Result.docx')


if __name__ == '__main__':
    try:
        main()
    except SystemExit as e:
        print('Error!', e)
        print('Press enter to exit (and fix the problem)')
        input()